from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
    elif level == 2:
        heading_style.font.size = Pt(15)
        heading_style.font.bold = True
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_title_page():
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Medliy')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('A Flutter-Based Telemedicine &\nHealthcare Consultation Platform')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('_' * 60)
    run.font.color.rgb = RGBColor(0, 102, 102)

    doc.add_paragraph()

    info_lines = [
        'A Project Report',
        'Submitted in Partial Fulfillment of the Requirements',
        'for the Degree of',
        '',
        'Bachelor of Technology',
        'in',
        'Computer Science & Engineering',
    ]
    for text in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(14)
        if text in ['A Project Report', 'Bachelor of Technology']:
            run.font.bold = True

    doc.add_paragraph()

    for _ in range(3):
        doc.add_paragraph()

    submitted = doc.add_paragraph()
    submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = submitted.add_run('Submitted By:')
    run.font.size = Pt(12)
    run.font.bold = True

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('[Student Name]')
    run.font.size = Pt(14)
    run.font.bold = True

    reg = doc.add_paragraph()
    reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = reg.add_run('Roll No: [Your Roll Number]')
    run.font.size = Pt(12)

    doc.add_paragraph()

    guide = doc.add_paragraph()
    guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = guide.add_run('Under the Guidance of:')
    run.font.size = Pt(12)
    run.font.bold = True

    guide_name = doc.add_paragraph()
    guide_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = guide_name.add_run('[Guide Name]\n[Designation]')
    run.font.size = Pt(13)

    doc.add_paragraph()

    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dept.add_run('Department of Computer Science & Engineering')
    run.font.size = Pt(13)
    run.font.bold = True

    college = doc.add_paragraph()
    college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = college.add_run('[College/University Name]\n[City, State]')
    run.font.size = Pt(12)

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('2025-2026')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_page_break()


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.75)
    return p


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def format_cell(cell, text, bold=False, color=None, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        format_cell(cell, header, bold=True, color=RGBColor(255, 255, 255), size=Pt(11))
        set_cell_shading(cell, '00796B')

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            format_cell(cell, str(cell_text), size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'E0F2F1')

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width

    doc.add_paragraph()
    return table


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0)
    return p


def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('~' * 50)
    run.font.color.rgb = RGBColor(180, 180, 180)


# ===================== BUILD DOCUMENT =====================

add_title_page()

# TABLE OF CONTENTS
add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'Abstract', '3'),
    ('2.', 'Introduction', '4'),
    ('3.', 'Problem Statement', '5'),
    ('4.', 'Objectives', '6'),
    ('5.', 'Literature Review & Existing Systems', '7'),
    ('6.', 'System Requirements', '8'),
    ('7.', 'System Architecture & Design', '10'),
    ('8.', 'Technology Stack', '13'),
    ('9.', 'Implementation', '15'),
    ('10.', 'Features', '18'),
    ('11.', 'Database Design', '20'),
    ('12.', 'User Interface', '22'),
    ('13.', 'Testing', '23'),
    ('14.', 'Limitations', '25'),
    ('15.', 'Future Scope', '26'),
    ('16.', 'Conclusion', '27'),
    ('17.', 'References', '28'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{num}  {title}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ============ 1. ABSTRACT ============
add_heading('1. Abstract', level=1)

add_para(
    'Medliy is a cross-platform telemedicine mobile application built with Flutter and Firebase '
    'that enables seamless communication between patients, doctors, and administrators. The platform '
    'allows patients to search for doctors by specialty, book appointments, purchase medicines from an '
    'integrated pharmacy, and communicate with doctors through real-time chat, voice messages, and '
    'voice/video calls.'
)
add_para(
    'Doctors can manage their schedules, handle appointments, prescribe medications, and consult with '
    'patients remotely. An admin panel allows the registration of verified doctors and management of '
    'the pharmacy inventory.'
)
add_para(
    'The application uses Agora Real-Time Communication SDK for voice and video calling, Firebase '
    'Authentication for secure login, Cloud Firestore for real-time data storage, and Firebase Storage '
    'for media file handling.'
)

doc.add_page_break()

# ============ 2. INTRODUCTION ============
add_heading('2. Introduction', level=1)

add_heading('2.1 Background', level=2)
add_para(
    'The healthcare industry has witnessed a rapid transformation with the advent of digital technology. '
    'Telemedicine - the remote delivery of healthcare services using telecommunications technology - has '
    'emerged as a vital solution for bridging the gap between patients and healthcare providers. The '
    'COVID-19 pandemic further accelerated the adoption of digital health solutions, making virtual '
    'consultations a necessity rather than a convenience.'
)

add_heading('2.2 About the Project', level=2)
add_para('Medliy (also referred to as "Medicore") is a comprehensive telemedicine platform designed to '
         'digitize the healthcare consultation process. The application provides a complete ecosystem where:')

add_bullet('Patients can discover doctors, book appointments, consult virtually, purchase medicines, and manage their health profiles.')
add_bullet('Doctors can manage their practice by setting availability, handling patient consultations, issuing prescriptions, and communicating with patients.')
add_bullet('Administrators can onboard verified doctors and manage the pharmacy inventory.')

add_para('The application is developed as a Flutter mobile application backed by Firebase cloud services, '
         'ensuring real-time data synchronization, secure authentication, and scalable infrastructure.')

doc.add_page_break()

# ============ 3. PROBLEM STATEMENT ============
add_heading('3. Problem Statement', level=1)

add_para('In many regions, accessing quality healthcare remains a significant challenge due to:')

add_bullet('Geographic barriers - Patients in rural or remote areas have limited access to specialists.')
add_bullet('Long wait times - Physical clinic visits involve long waiting periods for short consultations.')
add_bullet('Lack of transparency - Patients have no easy way to compare doctors, check availability, or read reviews.')
add_bullet('Pharmacy accessibility - Obtaining prescribed medicines often requires separate visits to pharmacies.')
add_bullet('Communication gaps - Post-consultation follow-ups and queries are difficult to manage through traditional channels.')

add_para(
    'There is a need for an integrated digital healthcare platform that combines appointment booking, '
    'virtual consultation, real-time communication, and pharmacy services into a single application.',
    bold=True
)

doc.add_page_break()

# ============ 4. OBJECTIVES ============
add_heading('4. Objectives', level=1)

add_para('The primary objectives of this project are:')

objectives = [
    'Develop a cross-platform mobile application using Flutter that works on Android, iOS, Web, macOS, and Windows.',
    'Implement role-based access control for three user types: Patient, Doctor, and Administrator.',
    'Enable doctor discovery and appointment booking with real-time availability scheduling.',
    'Integrate real-time communication including text chat, voice messages, image/document sharing, and voice/video calls.',
    'Build an integrated pharmacy with medicine catalog, shopping cart, and order management.',
    'Provide prescription management allowing doctors to issue prescriptions during consultations.',
    'Ensure secure authentication and data storage using Firebase services.',
    'Implement a doctor verification workflow where only admin-approved doctors can register on the platform.',
]

for i, obj in enumerate(objectives, 1):
    add_bullet(f'{obj}')

doc.add_page_break()

# ============ 5. LITERATURE REVIEW ============
add_heading('5. Literature Review & Existing Systems', level=1)

add_heading('5.1 Existing Telemedicine Platforms', level=2)

add_table(
    ['Platform', 'Features', 'Limitations'],
    [
        ['Practo', 'Doctor search, appointments, online consultation', 'Limited pharmacy integration, paid consultations'],
        ['Teladoc', 'Virtual visits, therapy, dermatology', 'US-only, insurance-dependent'],
        ['1mg', 'Pharmacy, lab tests, doctor consultation', 'Primarily pharmacy-focused'],
        ['MFine', 'AI-driven consultations, partner hospitals', 'Discontinued in 2023'],
        ['Amwell', 'Video visits, chronic care management', 'Enterprise-focused, not individual-friendly'],
    ]
)

add_heading('5.2 Research Findings', level=2)
add_bullet('According to WHO (2023), telemedicine can reduce healthcare costs by up to 30% and improve access in underserved areas.')
add_bullet('A study published in JMIR found that 76% of patients prefer telemedicine for follow-up consultations.')
add_bullet('Flutter has been recognized as a leading cross-platform framework, with over 500,000 apps published on both app stores (Google, 2024).')

add_heading('5.3 Gap Analysis', level=2)
add_para(
    'Most existing platforms are either region-specific, require insurance integration, or focus on a '
    'single aspect of healthcare. There is a need for an all-in-one platform that combines consultation, '
    'communication, and pharmacy services with a simple, accessible interface - which is what Medliy aims to provide.'
)

doc.add_page_break()

# ============ 6. SYSTEM REQUIREMENTS ============
add_heading('6. System Requirements', level=1)

add_heading('6.1 Hardware Requirements', level=2)
add_table(
    ['Component', 'Minimum Requirement'],
    [
        ['Developer Machine', '8 GB RAM, 4-core processor, 20 GB free disk space'],
        ['Android Device', 'Android 6.0 (API 23) or higher, 2 GB RAM'],
        ['iOS Device', 'iOS 12.0 or higher, 2 GB RAM'],
        ['Network', 'Broadband internet (development), 3G/4G/5G/Wi-Fi (usage)'],
    ],
    col_widths=[Cm(5), Cm(10)]
)

add_heading('6.2 Software Requirements', level=2)
add_table(
    ['Component', 'Specification'],
    [
        ['Operating System', 'Windows 10/11 (for Android development)'],
        ['Flutter SDK', '3.x (Dart SDK >=3.0.0 <4.0.0)'],
        ['IDE', 'Android Studio / VS Code (Latest stable)'],
        ['Firebase Account', 'For backend services configuration'],
        ['Agora Account', 'For real-time voice/video communication'],
    ],
    col_widths=[Cm(5), Cm(10)]
)

add_heading('6.3 Functional Requirements', level=2)
add_table(
    ['ID', 'Requirement', 'Priority'],
    [
        ['FR-01', 'User registration and login with role selection', 'High'],
        ['FR-02', 'Doctor search and filtering by specialty', 'High'],
        ['FR-03', 'Appointment booking with date and time slot selection', 'High'],
        ['FR-04', 'Real-time text messaging between patients and doctors', 'High'],
        ['FR-05', 'Voice and video calling between patients and doctors', 'High'],
        ['FR-06', 'Voice message recording and playback in chat', 'Medium'],
        ['FR-07', 'Image and document sharing in chat', 'Medium'],
        ['FR-08', 'Pharmacy with medicine catalog and shopping cart', 'High'],
        ['FR-09', 'Prescription creation by doctors', 'High'],
        ['FR-10', 'Doctor availability schedule management', 'High'],
        ['FR-11', 'Admin doctor registration and verification', 'High'],
        ['FR-12', 'Admin pharmacy inventory management', 'Medium'],
        ['FR-13', 'Doctor reviews and ratings', 'Medium'],
        ['FR-14', 'Online/offline status tracking', 'Low'],
        ['FR-15', 'Profile management with photo upload', 'Medium'],
    ]
)

add_heading('6.4 Non-Functional Requirements', level=2)
add_table(
    ['ID', 'Requirement', 'Description'],
    [
        ['NFR-01', 'Performance', 'App launch time < 3 seconds on modern devices'],
        ['NFR-02', 'Scalability', 'Firebase handles auto-scaling for database and storage'],
        ['NFR-03', 'Security', 'Firebase Authentication with role-based access control'],
        ['NFR-04', 'Availability', '99.9% uptime via Firebase cloud infrastructure'],
        ['NFR-05', 'Usability', 'Material Design 3 UI with consistent teal color theme'],
        ['NFR-06', 'Portability', 'Cross-platform support (Android, iOS, Web)'],
    ]
)

doc.add_page_break()

# ============ 7. SYSTEM ARCHITECTURE ============
add_heading('7. System Architecture & Design', level=1)

add_heading('7.1 High-Level Architecture', level=2)
add_para('The system follows a three-tier architecture with Client, Service, and Cloud layers:')

add_code_block(
    '┌─────────────────────────────────────────────────────────┐\n'
    '│                    CLIENT LAYER                         │\n'
    '│                                                         │\n'
    '│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐  │\n'
    '│  │   Patient     │  │   Doctor     │  │   Admin      │  │\n'
    '│  │   Mobile App  │  │   Mobile App │  │   Mobile App │  │\n'
    '│  └──────┬───────┘  └──────┬───────┘  └──────┬───────┘  │\n'
    '│         └────────┬────────┴──────────┬───────┘          │\n'
    '│         ┌────────▼───────────────────▼────────┐         │\n'
    '│         │      Flutter UI (Screens)           │         │\n'
    '│         └────────┬───────────────────┬────────┘         │\n'
    '│         ┌────────▼───────────────────▼────────┐         │\n'
    '│         │         SERVICE LAYER               │         │\n'
    '│         │  Auth | User | Doctor | Chat |      │         │\n'
    '│         │  Appointment | Pharmacy | Cart |    │         │\n'
    '│         │  Image | Call                       │         │\n'
    '│         └────────┬───────────────────┬────────┘         │\n'
    '├──────────────────┼───────────────────┼───────────────────┤\n'
    '│              CLOUD SERVICES LAYER                       │\n'
    '│    ┌─────────────▼───┐  ┌───▼──────────────┐          │\n'
    '│    │  Firebase Auth   │  │  Cloud Firestore  │          │\n'
    '│    └─────────────────┘  └──────────────────┘          │\n'
    '│    ┌─────────────────┐  ┌──────────────────┐          │\n'
    '│    │ Firebase Storage │  │  Agora RTC SDK   │          │\n'
    '│    └─────────────────┘  └──────────────────┘          │\n'
    '└─────────────────────────────────────────────────────────┘'
)

add_heading('7.2 Application Layer Architecture', level=2)
add_para('The application follows a Service-Oriented Architecture with clear separation of concerns:')

add_code_block(
    'lib/\n'
    '├── main.dart                    # Entry point, routes, theme\n'
    '├── firebase_options.dart        # Firebase configuration\n'
    '├── screens/                     # Presentation Layer (7 screens)\n'
    '│   ├── splash_screen.dart       # App launch, auto-login\n'
    '│   ├── auth_screen.dart         # Login/Register\n'
    '│   ├── patient_home_refactored  # Patient dashboard (4 tabs)\n'
    '│   ├── doctor_home.dart         # Doctor dashboard (4 tabs)\n'
    '│   ├── admin_home.dart          # Admin dashboard (3 tabs)\n'
    '│   ├── chat_screen.dart         # 1-to-1 messaging\n'
    '│   └── call_screen.dart         # Voice/video call\n'
    '├── services/                    # Business Logic (9 services)\n'
    '├── models/                      # Data Layer (9 models)\n'
    '├── components/                  # Reusable UI Widgets (3 files)\n'
    '└── utils/                       # Utilities (1 file)'
)

add_heading('7.3 Data Flow Diagram', level=2)

add_code_block(
    '    Patient                          Doctor\n'
    '       │                               │\n'
    '       │  1. Search Doctor             │\n'
    '       │──────────────────────►        │\n'
    '       │                               │\n'
    '       │  2. Book Appointment          │\n'
    '       │──────────────► Firestore ────► Doctor notified\n'
    '       │                               │\n'
    '       │  3. Chat / Voice / Video      │\n'
    '       │◄──────────────────────────────►│\n'
    '       │      (Real-time via Firestore  │\n'
    '       │       + Agora RTC)             │\n'
    '       │                               │\n'
    '       │  4. Buy Medicine              │\n'
    '       │──────────► Pharmacy ──► Order  │\n'
    '       │                               │\n'
    '       │  5. Rate Doctor               │\n'
    '       │──────────────────────► Firestore'
)

doc.add_page_break()

# ============ 8. TECHNOLOGY STACK ============
add_heading('8. Technology Stack', level=1)

add_heading('8.1 Frontend', level=2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Flutter', '3.x', 'Cross-platform UI framework'],
        ['Dart', '>=3.0.0 <4.0.0', 'Programming language'],
        ['Material Design 3', 'Built-in', 'UI component library'],
    ]
)

add_heading('8.2 Backend & Cloud', level=2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Firebase Auth', '5.0.0', 'Email/password authentication'],
        ['Cloud Firestore', '5.0.0', 'Real-time NoSQL database'],
        ['Firebase Storage', '12.0.0', 'Media file storage'],
        ['Firebase Core', '3.0.0', 'Firebase SDK initialization'],
    ]
)

add_heading('8.3 Communication', level=2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Agora RTC Engine', '6.5.4', 'Real-time voice and video calling'],
        ['Permission Handler', '11.4.0', 'Runtime permission management'],
    ]
)

add_heading('8.4 Supporting Libraries', level=2)
add_table(
    ['Library', 'Purpose'],
    [
        ['provider', 'State management'],
        ['image_picker', 'Camera and gallery image selection'],
        ['file_picker', 'Document file selection'],
        ['cached_network_image', 'Efficient image loading and caching'],
        ['record', 'Audio recording for voice messages'],
        ['shared_preferences', 'Local data persistence (cart)'],
        ['connectivity_plus', 'Network connectivity detection'],
        ['shimmer', 'Loading skeleton animations'],
        ['flutter_svg', 'SVG image rendering'],
        ['badges', 'Notification badge widgets'],
        ['uuid', 'Unique ID generation'],
        ['intl', 'Date and time formatting'],
        ['timeago', 'Human-readable time differences'],
        ['path_provider', 'File system directory access'],
        ['http', 'HTTP networking'],
        ['device_info_plus', 'Device information retrieval'],
    ]
)

doc.add_page_break()

# ============ 9. IMPLEMENTATION ============
add_heading('9. Implementation', level=1)

add_heading('9.1 Authentication System', level=2)
add_para(
    'The authentication system uses Firebase Authentication with email/password credentials. '
    'Three roles (Patient, Doctor, Admin) are managed through a single users collection in Firestore, '
    'differentiated by a role field.'
)

add_para('Doctor Registration Workflow:', bold=True)
add_code_block(
    'Admin adds doctor --> Pending record created -->\n'
    'Doctor registers --> Profile migrated -->\n'
    'Pending record deleted --> Doctor can login'
)

add_heading('9.2 Appointment System', level=2)
add_para('The appointment system follows a state machine workflow:')

add_code_block(
    '                ┌──────────┐\n'
    '                │  Booked  │  (Initial state)\n'
    '                └────┬─────┘\n'
    '                     │\n'
    '          ┌──────────┼──────────┐\n'
    '          ▼                     ▼\n'
    '    ┌───────────┐        ┌───────────┐\n'
    '    │ Confirmed │        │ Cancelled │\n'
    '    └─────┬─────┘        └───────────┘\n'
    '          │\n'
    '    ┌─────▼─────┐\n'
    '    │ Completed │\n'
    '    └───────────┘'
)

add_para('States: Pending -> Confirmed -> Completed | Cancelled', bold=True)

add_heading('9.3 Real-Time Chat System', level=2)
add_para(
    'The chat system uses Firestore real-time streams for instant message delivery. '
    'Chat rooms are created when a patient initiates a conversation with a doctor. '
    'Messages are stored as sub-documents under chats/{chatRoomId}/messages. '
    'Supports five message types: text, image, audio, document, call notification.'
)

add_heading('9.4 Voice & Video Calling', level=2)
add_para('The calling system uses Agora RTC SDK integrated with Firestore for call signaling:')

add_code_block(
    '1. Caller initiates call --> Firestore doc created (status: "ringing")\n'
    '2. Both callers join same Agora channel (named after chatRoomId)\n'
    '3. Remote user detected --> Call status updated to "active"\n'
    '4. Timer starts on connection\n'
    '5. Either party ends --> Status "ended" --> Agora channel left'
)

add_heading('9.5 Pharmacy System', level=2)
add_para(
    'The pharmacy system implements a basic e-commerce flow: Medicine catalog in Firestore, '
    'local cart persistence using SharedPreferences, checkout creating Firestore orders and '
    'decrementing stock.'
)

doc.add_page_break()

# ============ 10. FEATURES ============
add_heading('10. Features', level=1)

add_heading('10.1 Patient Features', level=2)
add_table(
    ['Feature', 'Description'],
    [
        ['Doctor Search', 'Search by name or specialty with real-time filtering'],
        ['Appointment Booking', 'Select date -> view available slots -> book appointment'],
        ['View Appointments', 'List all appointments with status badges and actions'],
        ['Pharmacy', 'Browse medicines, add to cart, adjust quantities, checkout'],
        ['Chat', 'Real-time messaging (text, images, documents, voice)'],
        ['Voice/Video Calls', 'Direct calls to doctors from the chat screen'],
        ['Profile Management', 'Edit name, phone, city, address, gender; upload photo'],
        ['Doctor Reviews', 'Rate and review doctors after consultation'],
    ]
)

add_heading('10.2 Doctor Features', level=2)
add_table(
    ['Feature', 'Description'],
    [
        ['Appointment Dashboard', 'Today\'s appointments, patient queue, recent activity'],
        ['Appointment Management', 'Confirm, complete, cancel appointments'],
        ['Prescriptions', 'Add text prescriptions to completed appointments'],
        ['Schedule Management', 'Set weekly availability with per-day time slots'],
        ['Chat', 'Real-time messaging with patients'],
        ['Voice/Video Calls', 'Direct calls to patients from the chat screen'],
        ['Profile Management', 'Edit profile, upload photo, view specialty'],
    ]
)

add_heading('10.3 Admin Features', level=2)
add_table(
    ['Feature', 'Description'],
    [
        ['Doctor Registration', 'Add new doctors with photo, specialty, availability'],
        ['Pharmacy Management', 'Add medicines with name, stock, price'],
        ['Inventory View', 'View complete pharmacy inventory'],
        ['Chat Access', 'View and participate in all conversations'],
    ]
)

doc.add_page_break()

# ============ 11. DATABASE DESIGN ============
add_heading('11. Database Design', level=1)

add_heading('11.1 Firestore Collections Structure', level=2)

add_para('The application uses six main Firestore collections:', bold=True)

add_code_block(
    'Firebase Firestore\n'
    '│\n'
    '├── users/                          # All users\n'
    '│   └── {userId}\n'
    '│       ├── uid, name, email, role\n'
    '│       ├── phoneNumber, profileImageUrl\n'
    '│       ├── gender, dateOfBirth, address, city\n'
    '│       ├── isOnline, lastSeen, createdAt\n'
    '│       ├── [Doctor fields]: specialty, experience,\n'
    '│       │   fee, availability, rating, totalReviews\n'
    '│       └── [Patient fields]: bloodGroup, allergies,\n'
    '│           medicalHistory\n'
    '│\n'
    '├── appointments/                   # All appointments\n'
    '│   └── {appointmentId}\n'
    '│       ├── patientId, doctorId, patientName, doctorName\n'
    '│       ├── appointmentDate, timeSlot, status\n'
    '│       ├── notes, prescription, consultationFee\n'
    '│       └── createdAt\n'
    '│\n'
    '├── chats/                          # Chat rooms\n'
    '│   └── {chatRoomId}\n'
    '│       ├── participants, participantNames\n'
    '│       ├── lastMessage, lastMessageTime, unreadCount\n'
    '│       └── messages/               # Sub-collection\n'
    '│           └── {messageId}\n'
    '│               ├── senderId, senderName, message\n'
    '│               ├── timestamp, isRead\n'
    '│               └── fileUrl, fileType\n'
    '│\n'
    '├── pharmacy/                       # Medicine inventory\n'
    '│   └── {medicineId}\n'
    '│       ├── name, stock, price\n'
    '│\n'
    '├── orders/                         # Pharmacy orders\n'
    '│   └── {orderId}\n'
    '│       ├── patientId, patientName, items[]\n'
    '│       └── totalAmount, createdAt\n'
    '│\n'
    '└── calls/                          # Call signaling\n'
    '    └── {callId}\n'
    '        ├── callerId, receiverId, type\n'
    '        ├── status (ringing|active|ended)\n'
    '        └── participants, startedAt'
)

add_heading('11.2 Entity Relationship Diagram', level=2)

add_code_block(
    '┌──────────┐     ┌──────────────┐     ┌──────────┐\n'
    '│  Users   │────<│ Appointments │>────│  Users   │\n'
    '│ (Patient)│     │              │     │ (Doctor) │\n'
    '└────┬─────┘     └──────────────┘     └──────────┘\n'
    '     │\n'
    '     │ 1:N\n'
    '     ▼\n'
    '┌──────────┐     ┌──────────────┐\n'
    '│  Chats   │────<│  Messages    │\n'
    '└────┬─────┘     └──────────────┘\n'
    '     │\n'
    '     │ 1:N\n'
    '     ▼\n'
    '┌──────────┐\n'
    '│  Calls   │\n'
    '└──────────┘\n\n'
    '┌──────────┐     ┌──────────────┐\n'
    '│ Pharmacy │────<│   Orders     │\n'
    '└──────────┘     └──────────────┘'
)

doc.add_page_break()

# ============ 12. USER INTERFACE ============
add_heading('12. User Interface', level=1)

add_heading('12.1 Design Principles', level=2)
add_bullet('Material Design 3 with a custom teal color theme (#00796B primary)')
add_bullet('Consistent border radius (12-16px) across all components')
add_bullet('Card-based layouts with white backgrounds and subtle shadows')
add_bullet('Shimmer loading for skeleton states during data loading')
add_bullet('Responsive layouts using MediaQuery for different screen sizes')

add_heading('12.2 Color Palette', level=2)
add_table(
    ['Color', 'Hex Code', 'Usage'],
    [
        ['Primary Teal', '#00796B', 'App bar, buttons, accents'],
        ['Dark Teal', '#004D40', 'Call screen background'],
        ['Indigo', '#1A237E', 'Chat bubbles (sent)'],
        ['Light Gray', '#F5F7FA', 'Screen backgrounds'],
        ['White', '#FFFFFF', 'Cards, input fields'],
    ]
)

add_heading('12.3 Screen Descriptions', level=2)
add_table(
    ['Screen', 'Layout'],
    [
        ['Splash', 'Gradient background, animated logo, auto-redirect'],
        ['Auth', 'Login/Register form with role tabs, forgot password'],
        ['Patient Home', 'Bottom nav (Home/Appointments/Pharmacy/Profile) - 4 tabs'],
        ['Doctor Home', 'Bottom nav (Appointments/Inbox/Schedule/Profile) - 4 tabs'],
        ['Admin Home', 'Bottom nav (Add Doctor/Pharmacy/Profile) - 3 tabs'],
        ['Chat', 'AppBar + message list + input bar'],
        ['Call', 'Full-screen with video views or avatar + controls'],
    ]
)

doc.add_page_break()

# ============ 13. TESTING ============
add_heading('13. Testing', level=1)

add_heading('13.1 Testing Approach', level=2)
add_table(
    ['Test Type', 'Method'],
    [
        ['Static Analysis', 'Flutter Analyzer for code quality and type safety'],
        ['Unit Testing', 'Service layer methods (planned)'],
        ['Widget Testing', 'Individual widget rendering (planned)'],
        ['Integration Testing', 'End-to-end user flows (planned)'],
        ['Manual Testing', 'Device testing on Android emulator and physical devices'],
    ]
)

add_heading('13.2 Tested Functionalities', level=2)
add_table(
    ['Feature', 'Status', 'Notes'],
    [
        ['User registration (Patient)', 'Passed', 'Email/password with role assignment'],
        ['User registration (Doctor)', 'Passed', 'Requires prior pending record'],
        ['Login/Logout', 'Passed', 'Role-based routing to correct home'],
        ['Doctor search', 'Passed', 'Real-time filtering by name/specialty'],
        ['Appointment booking', 'Passed', 'Date -> slot selection -> confirmation'],
        ['Appointment status flow', 'Passed', 'Pending -> Confirmed -> Completed'],
        ['Chat messaging', 'Passed', 'Real-time bidirectional messaging'],
        ['Voice message recording', 'Passed', 'Record -> upload -> send'],
        ['Image attachment', 'Passed', 'Camera/gallery pick -> upload -> display'],
        ['Document attachment', 'Passed', 'File picker -> upload -> display'],
        ['Voice/Video calling', 'Passed', 'Agora RTC with Firestore signaling'],
        ['Pharmacy browsing', 'Passed', 'Medicine catalog with search'],
        ['Shopping cart', 'Passed', 'Add/remove/quantity with persistence'],
        ['Checkout', 'Passed', 'Order creation and stock decrement'],
        ['Prescription management', 'Passed', 'Doctor adds, patient views'],
        ['Profile editing', 'Passed', 'Name, phone, city, address, photo'],
        ['Schedule management', 'Passed', 'Per-day toggle and time slot editing'],
        ['Admin doctor registration', 'Passed', 'Form with availability picker'],
        ['Admin pharmacy management', 'Passed', 'Add/view medicines'],
    ]
)

doc.add_page_break()

# ============ 14. LIMITATIONS ============
add_heading('14. Limitations', level=1)

limitations = [
    'No push notifications: Incoming calls and messages are only detected when the user is actively using the app. Firebase Cloud Messaging (FCM) is not yet integrated.',
    'No payment gateway: The pharmacy checkout is simulated without actual payment processing.',
    'Base64 profile images: Storing profile images as base64 in Firestore increases document size and may impact performance with many users.',
    'No offline support: The app requires an active internet connection for all operations. No local caching or offline-first architecture.',
    'No video call recording: Calls are real-time only with no recording or playback capability.',
    'No multi-language support: The application is currently English-only.',
    'No automated tests: Unit, widget, and integration tests are planned but not yet implemented.',
    'Single prescription format: Prescriptions are plain text only, without structured medication data.',
    'No appointment reminders: No notification system for upcoming appointments.',
    'No doctor verification documents: Doctor verification is a simple boolean flag without document upload.',
]

for i, lim in enumerate(limitations, 1):
    add_bullet(f'{lim}')

doc.add_page_break()

# ============ 15. FUTURE SCOPE ============
add_heading('15. Future Scope', level=1)

future = [
    'Push Notifications - Integrate Firebase Cloud Messaging (FCM) for incoming call alerts, message notifications, and appointment reminders.',
    'Payment Integration - Add Stripe, Razorpay, or similar payment gateway for pharmacy orders and consultation fee payments.',
    'AI-Powered Features - Implement symptom checker, medicine interaction alerts, and doctor recommendation engine using machine learning.',
    'Multi-Language Support - Add localization for Hindi, Tamil, and other regional languages to improve accessibility.',
    'Video Consultation Recording - Allow doctors to record consultations (with patient consent) for medical records.',
    'Health Records - Implement a digital health record system for medical reports, lab results, and vaccination records.',
    'Group Consultations - Support multi-participant video calls for group consultations or medical conferences.',
    'Pharmacy Delivery Tracking - Integrate with logistics APIs for real-time order tracking.',
    'Automated Test Suite - Implement comprehensive unit, widget, and integration tests.',
    'Offline Mode - Implement local database caching using Hive or Drift for offline access.',
    'Analytics Dashboard - Add analytics for doctors to view patient statistics and revenue reports.',
    'White-Label Solution - Refactor architecture to support multi-tenant deployments.',
]

for item in future:
    add_bullet(f'{item}')

doc.add_page_break()

# ============ 16. CONCLUSION ============
add_heading('16. Conclusion', level=1)

add_para(
    'Medliy is a comprehensive telemedicine platform that addresses the growing need for digital '
    'healthcare solutions. Built with Flutter and Firebase, it provides a cross-platform solution '
    'with real-time communication capabilities through text chat, voice messages, image/document '
    'sharing, and voice/video calling via Agora RTC.'
)
add_para(
    'The application successfully implements a multi-role ecosystem with Patient, Doctor, and Admin '
    'interfaces, each tailored to their specific needs. The patient can discover doctors, book '
    'appointments, purchase medicines, and consult virtually. The doctor can manage their schedule, '
    'handle consultations, and issue prescriptions. The admin can onboard verified doctors and manage '
    'pharmacy inventory.'
)
add_para(
    'The use of Flutter ensures a single codebase across Android, iOS, and web platforms, while '
    'Firebase provides a scalable, serverless backend with real-time synchronization. The Agora RTC '
    'integration brings WhatsApp-like calling capabilities directly into the chat interface.'
)
add_para(
    'The project demonstrates practical application of modern mobile development technologies and '
    'provides a solid foundation for further enhancement toward a production-ready telemedicine platform.',
    bold=True
)

doc.add_page_break()

# ============ 17. REFERENCES ============
add_heading('17. References', level=1)

references = [
    'Flutter Documentation. https://docs.flutter.dev/',
    'Firebase Documentation. https://firebase.google.com/docs',
    'Agora Real-Time Communication. https://docs.agora.io/en/',
    'Material Design 3. https://m3.material.io/',
    'Cloud Firestore Documentation. https://firebase.google.com/docs/firestore',
    'Firebase Authentication. https://firebase.google.com/docs/auth',
    'WHO Global Report on Digital Health, 2023.',
    'JMIR - Journal of Medical Internet Research. https://www.jmir.org/',
    'Dart Programming Language. https://dart.dev/',
    'permission_handler for Flutter. https://pub.dev/packages/permission_handler',
]

for i, ref in enumerate(references, 1):
    add_bullet(f'[{i}] {ref}')

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Medliy_Project_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
